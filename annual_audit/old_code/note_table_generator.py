
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def inject_three_column_tables(src_docx, dst_docx, start_tag, end_tag, balance_df):
    doc = Document(src_docx)
    paragraphs = doc.paragraphs

    # 找到标签范围
    start_idx = end_idx = -1
    for i, p in enumerate(paragraphs):
        if start_tag in p.text:
            start_idx = i
        if end_tag in p.text:
            end_idx = i
        if start_idx != -1 and end_idx != -1:
            break

    if start_idx == -1 or end_idx == -1 or start_idx >= end_idx:
        print("⚠️ 未找到有效的标签范围，插入表格失败")
        doc.save(dst_docx)
        return

    # 删除标签之间内容
    for i in range(end_idx - 1, start_idx, -1):
        p = paragraphs[i]._element
        p.getparent().remove(p)

    insert_para = paragraphs[start_idx]

    for _, row in balance_df.iterrows():
        name = str(row.iloc[0]).strip()
        end_val = row.get("期末数", 0)
        start_val = row.get("年初数", 0)

        if (not isinstance(end_val, (int, float))) or (not isinstance(start_val, (int, float))):
            continue
        if abs(end_val) < 1e-6 and abs(start_val) < 1e-6:
            continue

        # 插入标题段落 {{table_name}} 替换为 科目名称
        title = insert_para.insert_paragraph_before(name)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title.runs[0]
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "宋体"
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 插入表格
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.autofit = True

        hdr = table.rows[0].cells
        hdr[0].text = "科目"
        hdr[1].text = "期末数"
        hdr[2].text = "年初数"

        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = f"{end_val:,.2f}"
        row_cells[2].text = f"{start_val:,.2f}"

        # 表格前后插空段
        insert_para.insert_paragraph_before("")

    doc.save(dst_docx)
    print("✅ 附注美化表格已生成 →", dst_docx)
