import os
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

os.chdir(os.path.dirname(os.path.abspath(__file__)))

task_file = "task.xlsx"
mapping_file = "mappings.xlsx"
balance_file = "balan.xlsx"
activity_file = "yewu.xlsx"
main_template = "shenjishuoming.docx"
note_template = "fuzhu.docx"
output_main = "审计事项说明.docx"
output_note = "报表附注.docx"

def load_clean_df(path):
    df = pd.read_excel(path)
    df.columns = df.columns.str.strip()
    df.iloc[:, 0] = df.iloc[:, 0].astype(str).str.strip()
    return df

def generate_main_report():
    activity_df = load_clean_df(activity_file)
    balance_df = load_clean_df(balance_file)
    map_df = pd.read_excel(mapping_file)
    map_df['column'] = map_df['column'].fillna('本年累计数_合计').astype(str).str.strip()
    map_df['project_name'] = map_df['project_name'].astype(str).str.strip()
    map_df['source_sheet'] = map_df['source_sheet'].astype(str).str.strip()

    values = {}
    for _, row in map_df.iterrows():
        sheet = row['source_sheet']
        item = row['project_name']
        col = row['column']
        df = activity_df if sheet == "业务活动表" else balance_df if sheet == "资产负债表" else None
        val = df[df.iloc[:, 0] == item][col].values[0] if df is not None and col in df.columns and not df[df.iloc[:, 0] == item].empty else None
        values[(sheet, item)] = val

    context = {row['context_key']: values.get((row['source_sheet'], row['project_name']), None) for _, row in map_df.iterrows()}

    doc = DocxTemplate(main_template)

    task_df = pd.read_excel(task_file)
    if not task_df.empty:
        task_context = task_df.iloc[0].dropna().to_dict()
        context.update(task_context)

    doc.render(context)
    doc.save(output_main)
    print("✅ 审计事项说明生成完成")

def generate_note_report():
    activity_df = load_clean_df(activity_file)
    balance_df = load_clean_df(balance_file)
    combined_df = pd.concat([balance_df, activity_df], axis=0, ignore_index=True)

    map_df = pd.read_excel(mapping_file)
    map_df['column'] = map_df['column'].fillna('本年累计数_合计').astype(str).str.strip()
    map_df['project_name'] = map_df['project_name'].astype(str).str.strip()
    map_df['source_sheet'] = map_df['source_sheet'].astype(str).str.strip()

    values = {}
    for _, row in map_df.iterrows():
        sheet = row['source_sheet']
        item = row['project_name']
        col = row['column']
        df = activity_df if sheet == "业务活动表" else balance_df if sheet == "资产负债表" else None
        val = df[df.iloc[:, 0] == item][col].values[0] if df is not None and col in df.columns and not df[df.iloc[:, 0] == item].empty else None
        values[(sheet, item)] = val

    context = {row['context_key']: values.get((row['source_sheet'], row['project_name']), None) for _, row in map_df.iterrows()}

    task_df = pd.read_excel(task_file)
    if not task_df.empty:
        task_context = task_df.iloc[0].dropna().to_dict()
        context.update(task_context)

    # 渲染模板并保存为临时文件
    intermediate_docx = "temp_note_rendered.docx"
    doc_tpl = DocxTemplate(note_template)
    doc_tpl.render(context)
    doc_tpl.save(intermediate_docx)

    # 用 python-docx 重新加载渲染后的文档
    doc = Document(intermediate_docx)
    # 删除所有残留 {tableX_starts} 或 {tableX_ends} 标签段落
    for p in doc.paragraphs[:]:
        if p.text.strip().startswith("{table") and p.text.strip().endswith("_starts}"):
            p._element.getparent().remove(p._element)
        elif p.text.strip().startswith("{table") and p.text.strip().endswith("_ends}"):
            p._element.getparent().remove(p._element)
        elif p.text.strip() == "":  # 删除空段落
            p._element.getparent().remove(p._element)
    paras = doc.paragraphs

    start_tag = "{table1_starts}"
    end_tag = "{table5_ends}"
    start_idx = end_idx = -1
    for i, p in enumerate(paras):
        if start_tag in p.text:
            start_idx = i
        if end_tag in p.text:
            end_idx = i
        if start_idx != -1 and end_idx != -1:
            break

    if 0 <= start_idx <= end_idx:
        for i in range(end_idx, start_idx - 1, -1):
            if i < len(paras):
                paras[i]._element.getparent().remove(paras[i]._element)
        paras = doc.paragraphs

    insert_anchor = None
    if start_idx > 0 and start_idx - 1 < len(paras):
        insert_anchor = paras[start_idx - 1]._element
    else:
        insert_anchor = doc.element.body

    counter = 1
    for _, row in combined_df.iterrows():
        name = str(row.iloc[0]).strip()
        end_val = row.get("期末数", 0)
        start_val = row.get("年初数", 0)
        if not isinstance(end_val, (int, float)) or not isinstance(start_val, (int, float)):
            continue
        if abs(end_val) < 1e-6 and abs(start_val) < 1e-6:
            continue

        start_para = doc.add_paragraph(f"{{table{counter}_starts}}")
        start_para.runs[0].font.name = "宋体"
        start_para.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        insert_anchor.addnext(start_para._element)
        insert_anchor = start_para._element

        title_para = doc.add_paragraph(f"{counter}. {name}")
        title_para.paragraph_format.first_line_indent = Inches(0.74)
        run = title_para.runs[0]
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        insert_anchor.addnext(title_para._element)
        insert_anchor = title_para._element

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "科目"
        hdr[1].text = "期末数"
        hdr[2].text = "年初数"
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = f"{end_val:,.2f}"
        row_cells[2].text = f"{start_val:,.2f}"
        insert_anchor.addnext(table._element)
        insert_anchor = table._element

        end_para = doc.add_paragraph(f"{{table{counter}_ends}}")
        end_para.runs[0].font.name = "宋体"
        end_para.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        insert_anchor.addnext(end_para._element)
        insert_anchor = end_para._element

        empty_para = doc.add_paragraph()
        insert_anchor.addnext(empty_para._element)
        insert_anchor = empty_para._element

        counter += 1

    doc.save(output_note)
    print("✅ 报表附注生成完成")

if __name__ == "__main__":
    generate_main_report()
    generate_note_report()
    print("✅✅ 全部报告生成完毕！")
